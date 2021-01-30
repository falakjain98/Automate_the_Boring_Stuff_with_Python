#! python3
# customInvite.py - create customized invite documents using names stored in a txt file
# author: Falak Jain

import docx
from docx.shared import Pt

doc = docx.Document()
with open('guests.txt') as guests:
    for guest_name in guests.readlines():
        p = doc.add_paragraph()
        p.add_run('It would be a pleasure to have the company of').italic = True
        p.alignment = 1
        p.style.font.size = Pt(16)
        p = doc.add_paragraph()
        p.add_run(guest_name.rstrip('\n')).italic = True
        p.alignment = 1
        p = doc.add_paragraph()
        p.add_run('at 11010 Memory Lane on the Evening of').italic = True
        p.alignment = 1
        p = doc.add_paragraph()
        p.add_run('April 1st').italic = True
        p.alignment = 1
        p = doc.add_paragraph()
        p.add_run("at 7 o'clock").italic = True
        p.alignment = 1
        doc.add_page_break()
doc.save('invite.docx')

